from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("responsive-frontend-audit-and-plan.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5D6875"
GRID = "D5DEE8"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
RISK_FILL = "FFF3F0"
RISK = "9B1C1C"
SUCCESS_FILL = "EEF8F1"
SUCCESS = "176B3A"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_in) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    indent = table._tbl.tblPr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table._tbl.tblPr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    dxa_widths = [round(w * 1440) for w in widths_in]
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, dxa_widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, dxa_widths):
            cell.width = Inches(width / 1440)
            cell._tc.tcPr.tcW.set(qn("w:w"), str(width))
            cell._tc.tcPr.tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)


def set_paragraph_border(paragraph, color=BLUE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def set_run_font(run, size=11, color=INK, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("CLASSPHERE | RESPONSIVE FRONTEND AUDIT")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Internal planning document | July 2026")
    set_run_font(run, size=8.5, color=MUTED)


def add_para(doc, text: str = "", *, bold_prefix: str | None = None, color=INK, italic=False, after=6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color, italic=italic)


def add_bullets(doc, items: list[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_run_font(r)


def add_numbered(doc, items: list[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_run_font(r)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, label: str, text: str, *, risk=False) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, RISK_FILL if risk else CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, color=RISK if risk else DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=RISK if risk else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> None:
    doc = Document()
    style_doc(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("FRONTEND AUDIT & IMPLEMENTATION PLAN")
    set_run_font(r, size=9.5, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Responsive Classphere Web Application")
    set_run_font(r, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Mobile, tablet, desktop and test-taking experience | Audit baseline: July 18, 2026")
    set_run_font(r, size=12, color=MUTED)
    set_paragraph_border(p)

    add_table(doc, ["Scope", "Decision"], [
        ["Included", "Frontend layouts, navigation shells, test interface, states, accessibility and device QA."],
        ["Excluded", "Backend logic, data models, scoring, auth policy and API behavior. This plan changes presentation and client interaction only."],
        ["Primary outcome", "A single responsive web application: a full sidebar only on desktop, an accessible drawer at tablet/mobile sizes, and an isolated no-sidebar test environment."],
        ["Design constraint", "Keep the established NTA-style question palette and exam behavior. Improve hierarchy and adaptability; do not turn the test into a generic dashboard."],
    ], [1.875, 4.625])

    doc.add_heading("Executive finding", level=1)
    add_para(doc, "The application has a capable shared design system and many individual pages already contain responsive Tailwind utilities. However, responsiveness is not governed by one contract. The app shell transitions at 768px, while the intended tablet behavior requires the sidebar to be absent until desktop. The test route avoids the global shell, but its question palette is a desktop sidebar that falls into the page flow below 1280px rather than becoming a dedicated tablet/mobile navigator.")
    add_callout(doc, "Recommendation.", "Treat responsive behavior as a platform feature, not a page-by-page patch. Establish the breakpoints and shell patterns first; then migrate screens in priority order.")

    doc.add_heading("Audit evidence and priority findings", level=1)
    add_table(doc, ["Priority", "Finding", "Evidence", "Required direction"], [
        ["P0", "Sidebar appears across tablet widths", "AppShell and Sidebar use md (768px) to switch from MobileNav to a 280px desktop sidebar.", "Move the desktop shell breakpoint to lg (1024px). Use the drawer/top bar below it."],
        ["P0", "Superadmin bypasses the responsive shell", "The superadmin layout renders Sidebar directly and never renders MobileNav.", "Reuse the shared responsive shell or create a shared admin shell with the same 1024px contract."],
        ["P0", "Test palette becomes inline content on tablet", "The test grid creates the right column only at xl (1280px); below that, QuestionNavigator is rendered after the question card.", "Use a test-only drawer/sheet on tablet and mobile; keep a fixed right palette only on desktop."],
        ["P0", "Test loading reveals implementation language", "The live test loading state says 'Loading questions from backend'.", "Retain the skeleton and use no copy or student-facing wording such as 'Preparing your test'."],
        ["P1", "Mobile navigation has a divergent menu", "MobileNav omits some routes present in Sidebar, including institute Tests and teacher Batch/Doubts.", "Generate both desktop and mobile navigation from one role-aware source of truth."],
        ["P1", "Some product screens remain desktop-first", "Examples include fixed-width chat, profile side navigation, superadmin table/list pages and 24px desktop gutters on narrow devices.", "Migrate with shared responsive page, filter, table and panel primitives."],
        ["P1", "No consistent overflow contract", "Question markdown protects images and display math, but tables and several fixed-width panels can exceed narrow viewports.", "Define component-level horizontal-scroll and max-width rules rather than relying on page-level overflow clipping."],
        ["P2", "Touch/accessibility behavior is incomplete", "Mobile drawer locks scroll and closes on route change, but lacks Escape handling, focus containment and explicit menu semantics.", "Use a dialog-grade drawer/sheet primitive with focus return, Escape close and labelled controls."],
    ], [0.72, 1.44, 2.12, 2.22])

    doc.add_heading("Responsive product contract", level=1)
    add_para(doc, "The following breakpoints are a product contract, not merely CSS values. They remove ambiguity when new screens are added.")
    add_table(doc, ["Range", "Mode", "Navigation", "Page behavior"], [
        ["0-639px", "Phone", "Compact top bar + modal drawer. Never show the permanent sidebar.", "One-column content; 16px page gutters; action groups stack or use a safe two-column grid."],
        ["640-1023px", "Tablet", "Same compact top bar + drawer. This is deliberately still non-sidebar.", "Two-column cards where useful; tables adapt into cards or a controlled horizontal region."],
        ["1024-1279px", "Desktop", "Permanent 280px sidebar. No mobile top bar.", "Dashboard/workspace layout; tests may keep a desktop palette if there is sufficient room."],
        ["1280px+", "Wide desktop", "Permanent sidebar.", "Optional wider grid density and the full right-side NTA question palette."],
    ], [1.05, 1.15, 2.15, 2.15])
    add_callout(doc, "Breakpoint rule.", "Use lg as the desktop navigation boundary. Avoid introducing one-off breakpoints unless a component has an explicit tested reason.")

    doc.add_heading("Target shell architecture", level=1)
    doc.add_heading("1. Standard application shell", level=2)
    add_bullets(doc, [
        "One role-aware navigation configuration feeds Sidebar and MobileNav, eliminating missing or inconsistent menu routes.",
        "Desktop sidebar is visible only at lg and above. At all smaller widths, render a 64px top bar with an accessible menu trigger and drawer.",
        "Content container owns vertical scrolling. Pages use a shared responsive gutter token: 16px phone, 24px tablet, 32px desktop where appropriate.",
        "Drawer closes on navigation, backdrop click and Escape; it traps focus while open and returns focus to the menu trigger when closed.",
        "Do not use global overflow clipping as the fix for layout defects. Components that genuinely need a horizontal scroller must expose it intentionally and visibly.",
    ])
    doc.add_heading("2. Exam/test shell", level=2)
    add_bullets(doc, [
        "Live tests remain outside the dashboard shell: no application sidebar, no dashboard mobile nav, no competing product navigation.",
        "A compact exam header keeps the timer, identity and submit affordance visible without consuming excessive phone height.",
        "Desktop: question card plus right NTA palette. Tablet: question card plus a palette trigger that opens a right drawer. Phone: palette trigger opens a full-height sheet with sections and status legend.",
        "Question actions use two columns on phones and four columns only when there is room; every action remains at least 44px high and preserves the current NTA labels.",
        "Question images, option images, markdown tables and display math must be bounded by the content width. Long formulae get an intentional in-place horizontal scroller, never page-level overflow.",
        "Loading uses the existing skeleton structure and a neutral student-facing message, not backend terminology.",
    ])

    doc.add_heading("Screen-family audit", level=1)
    add_table(doc, ["Family", "Current state", "Implementation objective"], [
        ["Student dashboard, tests, DPPs, results", "Generally responsive in cards/grids, but tabs, filters and activity surfaces need a shared narrow-width pattern.", "Prioritize the student journey: tests list, launch card, live test, results and revision actions."],
        ["Live test", "Strong content rendering baseline; layout has no dedicated tablet/mobile palette behavior.", "Deliver first as a protected test-mode shell, without changing question/attempt logic."],
        ["Teacher/institute workspace", "Mixed. Several list pages adapt well, while forms use desktop gutters and long control rows.", "Convert headers, filters, tables and bulk actions to reusable responsive primitives."],
        ["Superadmin", "Own layout bypasses mobile navigation; numerous fixed-width lists and desktop paddings.", "Bring it under the same shell contract and render dense data as mobile cards or explicit scroll regions."],
        ["Profile/settings/help/doubts", "Profile and doubts include fixed desktop panels; settings has a workable horizontal tab fallback.", "Make conversational tools and profile navigation single-column on phone, two-pane only above tablet."],
        ["Auth, empty/error/loading states", "Mostly centered and safe, but state language and padding are inconsistent.", "Standardize skeletons and student-facing status copy; preserve 16px input text to prevent iOS zoom."],
    ], [1.45, 2.55, 2.45])

    doc.add_heading("Implementation sequence", level=1)
    add_numbered(doc, [
        "Create a responsive token layer: shell breakpoint, page gutters, safe-area spacing, touch target size, content max widths and intentional overflow helpers.",
        "Refactor AppShell, Sidebar and MobileNav into a single responsive navigation system. Move the breakpoint from md to lg and centralize role menus.",
        "Refactor the superadmin layout to use that same shell behavior. Verify auth guard behavior remains unchanged.",
        "Implement test-mode responsiveness: compact header, palette drawer/sheet, desktop palette rules, responsive action bar and non-technical loading skeleton wording.",
        "Migrate the student journey screens: dashboard, Tests Hub, assignments/DPPs, results, analytics and revision/mistake diary.",
        "Migrate institute, teacher and superadmin dense-data screens using shared filter bars, responsive tables/cards and mobile-safe forms.",
        "Run the device, accessibility and regression matrix before treating the work as complete.",
    ])
    add_callout(doc, "Scope safeguard.", "Each phase is presentation-only. No API route, request shape, auth rule, scoring rule, database migration or backend service should be changed for this initiative.")

    doc.add_heading("Acceptance criteria", level=1)
    add_table(doc, ["Area", "Done means"], [
        ["Navigation", "No permanent sidebar below 1024px. Every desktop route is reachable from the compact navigation for the same role. Drawer is keyboard and screen-reader operable."],
        ["Live test", "No global navigation/sidebar. Timer and submit remain usable at 320px. Question palette works as a drawer/sheet below desktop and preserves all status meanings."],
        ["Content", "No accidental horizontal page scroll at 320, 375, 390, 414, 768, 820, 1024, 1280 and 1440 CSS pixels. Images, equations and tables have defined containment."],
        ["Controls", "Interactive targets are at least 44x44px where practical. Input text is 16px or greater on phone. Focus is visible, order is logical and modals/drawers restore focus."],
        ["States", "Loading, error and empty states fit narrow screens and never expose internal implementation language such as 'backend'."],
        ["Visual quality", "Existing Core 2.0 style tokens remain consistent; changes do not reintroduce the previously rejected cluttered test palette redesign."],
    ], [1.25, 5.25])

    doc.add_heading("QA matrix", level=1)
    add_table(doc, ["Viewport", "Representative validation"], [
        ["320 / 375 / 390 / 414", "Student tests, live test, test palette sheet, submit confirmation, numerical input, long question and image question."],
        ["768 / 820", "Drawer navigation, dashboard cards, test palette drawer, filters, teacher/institute forms and list density."],
        ["1024 / 1280 / 1440", "Sidebar transition, test right palette, superadmin tables, cards and page gutters."],
        ["Keyboard + screen reader", "Menu, drawer, test palette, modal, submit flow, tab order, Escape and focus restoration."],
        ["Network/low-end device", "Skeleton stability, no layout shift during test load, images stay constrained and no expensive viewport-resize loops."],
    ], [1.45, 5.05])

    doc.add_heading("First implementation ticket", level=1)
    add_para(doc, "Start with the shared shell and live test mode. This produces the highest student-visible improvement while giving all other screens a reliable responsive foundation. The initial pull request should include only the navigation breakpoint, unified role menus, test header/palette adaptation, loading-copy change and automated viewport checks. Page-by-page migrations follow after the shell is accepted.")
    add_callout(doc, "Ready to proceed.", "The audit is complete and this document is the implementation baseline. The next task can begin with Phase 1 without revisiting backend logic.")

    doc.core_properties.title = "Responsive Classphere Web Application - Frontend Audit and Plan"
    doc.core_properties.subject = "Responsive frontend implementation plan"
    doc.core_properties.author = "Classphere Engineering"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
