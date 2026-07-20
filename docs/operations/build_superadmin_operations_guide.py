from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name("superadmin-operations-guide.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    table_pr.append(tbl_ind)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font(run, size=11, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(doc, text, *, bold=False, color=None, after=6, before=0, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text), size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text), size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 7)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), size=16 if level == 1 else 13, color=BLUE if level == 1 else DARK_BLUE, bold=True)
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [1.45, 1.25, 3.8])
    headers = ["Area", "Current state", "Operator action"]
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(value), size=10, color=DARK_BLUE, bold=True)
    for area, state, action in rows:
        cells = table.add_row().cells
        for idx, value in enumerate((area, state, action)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(value), size=9.5, bold=(idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("Classphere | Superadmin Operations Guide"), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("Internal operating guide - update after each infrastructure change"), size=8.5, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    font(title.add_run("SUPERADMIN OPERATIONS GUIDE"), size=23, color=DARK_BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    font(subtitle.add_run("What is live today, what is intentionally unavailable, and how to activate trustworthy operational data."), size=13, color=GRAY)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    set_table_widths(meta, [1.2, 5.3])
    for row, (label, value) in zip(meta.rows, [
        ("Audience", "Platform owner, engineering lead, and operations team"),
        ("Billing mode", "Free trials only. No invoice, charge, MRR, ARPU, churn, or payment reporting is enabled."),
        ("Access rule", "An institute must be active and have a valid trial or active paid entitlement before protected APIs are available."),
    ]):
        set_cell_shading(row.cells[0], "F2F4F7")
        for i, text in enumerate((label, value)):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(text), size=10.5, bold=(i == 0))

    add_heading(doc, "1. Operating principle")
    add_text(doc, "The panel must never turn unavailable telemetry into a plausible number. A card is either backed by an identified source of truth, explicitly marked unavailable, or not shown. This guide is the release checklist for moving an unavailable item into live operation.")
    add_bullet(doc, "Live means: the data is persisted, the calculation is defined, the source is monitored, and an owner is accountable.")
    add_bullet(doc, "Unavailable means: no synthetic values, no assumed percentages, and no operational decision should be made from the panel.")
    add_bullet(doc, "Billing remains deliberately disabled. The entitlement service supports paid plans later without opening access by default.")

    add_heading(doc, "2. Current source-of-truth map")
    add_status_table(doc, [
        ("Institutes and students", "Live", "Read from institutes, users, batches, and batch_students. Use server-side pagination once the tenant count grows."),
        ("Question bank", "Live", "Global uploads are atomic through the database RPC. Paper metadata edits and bulk actions use dedicated global-paper APIs."),
        ("Support and audit logs", "Live", "Support tickets are persisted. Privileged question-bank and institute actions are written to audit_logs."),
        ("Trial entitlement", "Live after migration 15", "Auth middleware checks the institute status plus a valid trial or active plan on every protected request."),
        ("Platform uptime", "Unavailable", "The API can report that it served the current request, but availability percentage requires an external monitor and a retained history."),
        ("R2 storage and CDN health", "Unavailable", "Configuration presence is not a health check. Add provider probes and provider metrics before showing health values."),
        ("AI usage and completion rate", "Unavailable", "Add durable usage and lifecycle events before enabling these analytics."),
        ("Revenue analytics", "Unavailable", "Enable only after a payment provider, invoice lifecycle, and subscription ledger are in place."),
        ("Runtime configuration", "Unavailable", "The configuration UI is intentionally non-operative until settings have typed runtime consumers and a safe rollout mechanism."),
    ])

    add_heading(doc, "3. Trial-only entitlement workflow")
    add_text(doc, "This is the intended behavior during the free-trial phase. It prevents a tenant from retaining API access just because its institute row exists.")
    add_bullet(doc, "Provision an institute with a trial duration from 1 to 24 months. The service creates a subscription record with plan_tier=trial and status=trialing.")
    add_bullet(doc, "Authentication verifies: user profile -> active institute -> valid entitlement. Expired trials return SUBSCRIPTION_REQUIRED. Suspended institutes return INSTITUTE_SUSPENDED.")
    add_bullet(doc, "Do not create trials lazily from a read endpoint. Every entitlement must originate from the provisioning or future billing workflow.")
    add_bullet(doc, "Before billing launches, a paid plan can be granted manually only by creating an active subscription record with a non-free plan_tier. Record the operator action in audit_logs.")
    add_text(doc, "Migration requirement: run 15_superadmin_entitlements_and_global_content.sql after migrations 10 and 14. Validate with one active trial, one expired trial, one suspended institute, and one global user with no institute.", bold=True, color=DARK_BLUE, after=10)

    add_heading(doc, "4. How to activate trustworthy uptime and health")
    add_text(doc, "Current API reachability is not uptime. Use an external synthetic monitor so outages can be detected even when the API itself is unavailable.")
    add_bullet(doc, "Create a public, dependency-aware health endpoint. It should expose separate API, database, Redis/queue, and worker readiness states; never return secrets or tenant data.")
    add_bullet(doc, "Probe the endpoint from at least three regions every 60 seconds using Better Uptime, UptimeRobot, Grafana Cloud, Datadog, or Cloudflare Health Checks.")
    add_bullet(doc, "Send probe results to an observability store. Calculate 7-day and 30-day uptime from the stored probe history, not from process uptime.")
    add_bullet(doc, "For R2, emit upload success/failure/latency counters and run a safe bucket List/Head probe with a dedicated health object. For CDN, use Cloudflare Analytics or edge request metrics.")
    add_bullet(doc, "Expose an internal telemetry aggregation API that reads those metrics. The superadmin UI should show Unknown when the monitor is silent or stale.")

    add_heading(doc, "5. How to activate AI and learning analytics")
    add_text(doc, "AI and completion KPIs require events, not estimates. Create a platform_usage_events table or send equivalent events to an analytics warehouse.")
    add_status_table(doc, [
        ("AI token usage", "Record provider, model, prompt_tokens, completion_tokens, cost, institute_id, user_id, feature, timestamp.", "Aggregate by month and institute. Retain raw events for reconciliation."),
        ("Completion rate", "Record attempt_started, attempt_submitted, attempt_expired with paper_id, institute_id, and timestamps.", "Define denominator explicitly: started attempts within a period, excluding invalidated attempts."),
        ("Top institutes", "Use distinct active student enrollment and scoped submitted attempts.", "Rank on an explicit period and display the selected metric in the label."),
        ("Booster/generated tests", "Record a test_generated event only after the generated paper is persisted.", "Count successful persisted items; never count requests that failed or were discarded."),
    ])

    add_heading(doc, "6. Future billing activation checklist")
    add_text(doc, "Do not turn on revenue cards merely because invoices can be inserted manually. Billing needs an independently reconcilable ledger.")
    add_bullet(doc, "Select a provider (for example Razorpay) and verify webhook signatures. Store provider event IDs with a unique constraint for idempotency.")
    add_bullet(doc, "Create subscription, invoice, payment, refund, and entitlement records. Keep monetary amounts as integer paise, not floating-point values.")
    add_bullet(doc, "Update entitlement only from verified payment/subscription events. Keep a grace-period policy explicit and test expiry, refund, cancellation, and duplicate webhooks.")
    add_bullet(doc, "Define MRR, collected revenue, ARPU, churn, and YTD revenue in writing. Each dashboard value must link to one query and a fixed reporting period.")
    add_bullet(doc, "Enable CSV export only after it includes a permission check, date filters, pagination, audit logging, and a safe download format."),

    add_heading(doc, "7. Runtime configuration activation checklist")
    add_text(doc, "Feature flags and infrastructure limits should remain disabled until they change real runtime behavior safely.")
    add_bullet(doc, "Define a typed setting schema with allowed values, defaults, validation, owner, rollout scope, and rollback instructions.")
    add_bullet(doc, "Use a cached server-side configuration provider with a short TTL and an explicit invalidation path. Do not query system_settings on every request.")
    add_bullet(doc, "Implement maintenance mode in middleware plus a safe allowlist for superadmins, health checks, and logout. Add a clear maintenance page and audit every change.")
    add_bullet(doc, "Implement each limit where it is enforced: upload size in multer, session timeout in auth/session expiry, concurrency at the edge or queue, and OMR rate in the job queue.")

    add_heading(doc, "8. Release verification")
    add_bullet(doc, "Run the database migrations, then run API and web production builds.")
    add_bullet(doc, "Create an institute with a one-month trial; confirm its staff and students can use protected APIs.")
    add_bullet(doc, "Expire its subscription in a test environment; confirm protected APIs return SUBSCRIPTION_REQUIRED while the superadmin can still manage the tenant.")
    add_bullet(doc, "Upload a valid global question paper, then force a failing upload in staging; confirm the failure creates neither a paper nor orphan question records.")
    add_bullet(doc, "Edit and bulk-edit global papers; confirm only public-practice papers are selectable and audit entries are produced.")
    add_bullet(doc, "Confirm analytics, revenue, uptime, storage, CDN, and AI cards show unavailable language until their corresponding data pipelines are complete.")

    doc.core_properties.title = "Classphere Superadmin Operations Guide"
    doc.core_properties.subject = "Operational activation and entitlement guide"
    doc.core_properties.author = "Classphere Engineering"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
